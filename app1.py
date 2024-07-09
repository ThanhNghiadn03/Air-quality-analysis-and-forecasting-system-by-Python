from flask import Flask, jsonify, render_template, request, send_file
import requests
import matplotlib
matplotlib.use('Agg')
from matplotlib import pyplot
import pandas as pd
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

@app.route('/')
def map_view():
    return render_template('map.html')

@app.route('/air_quality')
def get_air_quality():
    position = request.args.get('position', default='VietNam', type=str).lower()
    url = f"https://api.waqi.info/feed/{position}/?token=cdcfc9b6df90b585e02ec9d367c88ea755aca36c"
    response = requests.get(url)
    data = response.json()
    if data['status'] == "error" or data['status'] == "nope":
        return render_template('fail.html')
    return get_detail(data,position)

def check_path_exists(path):
    return os.path.exists(path)

def get_current_datetime():
    now = datetime.now()
    return now.strftime("%d-%m-%Y-%H")

def create_charts(data, location):
    # Tạo thư mục nếu chưa tồn tại
    os.makedirs(f'static/images/{location}', exist_ok=True)

    # Tạo DataFrame cho forecast data
    forecast_pm25 = pd.DataFrame(data['data']['forecast']['daily']['pm25'])
    forecast_pm10 = pd.DataFrame(data['data']['forecast']['daily']['pm10'])
    forecast_o3 = pd.DataFrame(data['data']['forecast']['daily']['o3'])

    # Tạo DataFrame cho IAQI data
    iaqi_data = pd.DataFrame.from_dict(data['data']['iaqi'], orient='index').reset_index()
    iaqi_data.columns = ['Pollutant', 'Value']

    # Vẽ biểu đồ bar chart cho IAQI, bao gồm giá trị âm
    pyplot.figure(figsize=(20, 16))
    pyplot.bar(iaqi_data['Pollutant'], iaqi_data['Value'], color='skyblue')
    pyplot.xlabel('Pollutants')
    pyplot.ylabel('Values')
    pyplot.title('IAQI Values for Different Pollutants (Including Negative Values)')
    pyplot.savefig(f'static/images/{location}/iaqi_bar.png')
    pyplot.close()

    # Vẽ biểu đồ bar chart cho forecast PM2.5
    pyplot.figure(figsize=(20, 16))
    pyplot.bar(forecast_pm25['day'], forecast_pm25['avg'], color='lightgreen')
    pyplot.xlabel('Day')
    pyplot.ylabel('PM2.5 Avg')
    pyplot.title('PM2.5 Forecast')
    pyplot.xticks(rotation=45)
    pyplot.savefig(f'static/images/{location}/pm25_forecast.png')
    pyplot.close()

    # Vẽ biểu đồ bar chart cho forecast PM10
    pyplot.figure(figsize=(20, 16))
    pyplot.bar(forecast_pm10['day'], forecast_pm10['avg'], color='orange')
    pyplot.xlabel('Day')
    pyplot.ylabel('PM10 Avg')
    pyplot.title('PM10 Forecast')
    pyplot.xticks(rotation=45)
    pyplot.savefig(f'static/images/{location}/pm10_forecast.png')
    pyplot.close()

    # Vẽ biểu đồ bar chart cho forecast O3
    pyplot.figure(figsize=(20, 16))
    pyplot.bar(forecast_o3['day'], forecast_o3['avg'], color='purple')
    pyplot.xlabel('Day')
    pyplot.ylabel('O3 Avg')
    pyplot.title('O3 Forecast')
    pyplot.xticks(rotation=45)
    pyplot.savefig(f'static/images/{location}/o3_forecast.png')
    pyplot.close()

    # Vẽ biểu đồ cột nhóm cho forecast PM2.5
    pyplot.figure(figsize=(20, 16))
    x = range(len(forecast_pm25))
    pyplot.bar([i - 0.2 for i in x], forecast_pm25['avg'], width=0.2, label='Avg', color='lightgreen', align='center')
    pyplot.bar(x, forecast_pm25['min'], width=0.2, label='Min', color='blue', align='center')
    pyplot.bar([i + 0.2 for i in x], forecast_pm25['max'], width=0.2, label='Max', color='red', align='center')
    pyplot.xlabel('Day')
    pyplot.ylabel('PM2.5 Values')
    pyplot.title('PM2.5 Forecast (Avg, Min, Max)')
    pyplot.xticks(ticks=x, labels=forecast_pm25['day'], rotation=45)
    pyplot.legend()
    pyplot.savefig(f'static/images/{location}/pm25_grouped.png')
    pyplot.close()

    # Vẽ biểu đồ cột nhóm cho forecast PM10
    pyplot.figure(figsize=(20, 16))
    x = range(len(forecast_pm10))
    pyplot.bar([i - 0.2 for i in x], forecast_pm10['avg'], width=0.2, label='Avg', color='orange', align='center')
    pyplot.bar(x, forecast_pm10['min'], width=0.2, label='Min', color='blue', align='center')
    pyplot.bar([i + 0.2 for i in x], forecast_pm10['max'], width=0.2, label='Max', color='red', align='center')
    pyplot.xlabel('Day')
    pyplot.ylabel('PM10 Values')
    pyplot.title('PM10 Forecast (Avg, Min, Max)')
    pyplot.xticks(ticks=x, labels=forecast_pm10['day'], rotation=45)
    pyplot.legend()
    pyplot.savefig(f'static/images/{location}/pm10_grouped.png')
    pyplot.close()

    # Vẽ biểu đồ cột nhóm cho forecast O3
    pyplot.figure(figsize=(20, 16))
    x = range(len(forecast_o3))
    pyplot.bar([i - 0.2 for i in x], forecast_o3['avg'], width=0.2, label='Avg', color='purple', align='center')
    pyplot.bar(x, forecast_o3['min'], width=0.2, label='Min', color='blue', align='center')
    pyplot.bar([i + 0.2 for i in x], forecast_o3['max'], width=0.2, label='Max', color='red', align='center')
    pyplot.xlabel('Day')
    pyplot.ylabel('O3 Values')
    pyplot.title('O3 Forecast (Avg, Min, Max)')
    pyplot.xticks(ticks=x, labels=forecast_o3['day'], rotation=45)
    pyplot.legend()
    pyplot.savefig(f'static/images/{location}/o3_grouped.png')
    pyplot.close()

def evaluate_aqi(aqi):
    if 0 <= aqi <= 50:
        return ("Good", "Air quality is considered satisfactory, and air pollution poses little or no risk.",
                "None")
    elif 51 <= aqi <= 100:
        return ("Moderate", "Air quality is acceptable; however, for some pollutants there may be a moderate health concern for a very small number of people who are unusually sensitive to air pollution.",
                "Active children and adults, and people with respiratory disease, such as asthma, should limit prolonged outdoor exertion.")
    elif 101 <= aqi <= 150:
        return ("Unhealthy for Sensitive Groups", "Members of sensitive groups may experience health effects. The general public is not likely to be affected.",
                "Active children and adults, and people with respiratory disease, such as asthma, should limit prolonged outdoor exertion.")
    elif 151 <= aqi <= 200:
        return ("Unhealthy", "Everyone may begin to experience health effects; members of sensitive groups may experience more serious health effects.",
                "Active children and adults, and people with respiratory disease, such as asthma, should avoid prolonged outdoor exertion; everyone else, especially children, should limit prolonged outdoor exertion.")
    elif 201 <= aqi <= 300:
        return ("Very Unhealthy", "Health warnings of emergency conditions. The entire population is more likely to be affected.",
                "Active children and adults, and people with respiratory disease, such as asthma, should avoid all outdoor exertion; everyone else, especially children, should limit outdoor exertion.")
    else:
        return ("Hazardous", "Health alert: everyone may experience more serious health effects.",
                "Everyone should avoid all outdoor exertion.")


@app.route('/detail')
def get_detail(data,position):
    location = data['data']['city']['name'].replace(" ", "").lower() + "_" + get_current_datetime().replace(":", "-")
    create_charts(data, location)
    aqi = data['data']['aqi']
    dominentpol = data['data']['dominentpol']
    # city_name = data['data']['city']['name']
    aqi_level, aqi_message, cautionary_statement = evaluate_aqi(aqi)
    return render_template('index.html', location=location, position=position, aqi=aqi, dominentpol=dominentpol, city_name=position, aqi_level=aqi_level, aqi_message=aqi_message, cautionary_statement=cautionary_statement)


@app.route('/export_word')
def export_word():
    location = request.args.get('location')
    position = request.args.get('position')
    aqi = request.args.get('aqi')
    dominentpol = request.args.get('dominentpol')
    city_name = request.args.get('city_name')
    aqi_level = request.args.get('aqi_level')
    aqi_message = request.args.get('aqi_message')
    cautionary_statement = request.args.get('cautionary_statement')
    
    document = Document()
    document.add_heading('Air Quality Report', 0)

    document.add_heading('City:', level=1)
    document.add_paragraph(city_name)

    document.add_heading('Current AQI:', level=1)
    document.add_paragraph(f'{aqi} ({aqi_level})')
    document.add_paragraph(aqi_message)

    document.add_heading('Cautionary Statement:', level=1)
    document.add_paragraph(cautionary_statement)

    document.add_heading('Dominent Pollutant:', level=1)
    document.add_paragraph(dominentpol)

    document.add_heading('Charts:', level=1)
    
    document.add_heading('IAQI Bar Chart', level=2)
    document.add_picture(f'static/images/{location}/iaqi_bar.png', width=Inches(6))

    document.add_heading('PM2.5 Forecast', level=2)
    document.add_picture(f'static/images/{location}/pm25_forecast.png', width=Inches(6))
    
    document.add_heading('PM2.5 Forecast (Grouped Bar Chart)', level=2)
    document.add_picture(f'static/images/{location}/pm25_grouped.png', width=Inches(6))

    document.add_heading('PM10 Forecast', level=2)
    document.add_picture(f'static/images/{location}/pm10_forecast.png', width=Inches(6))

    document.add_heading('PM10 Forecast (Grouped Bar Chart)', level=2)
    document.add_picture(f'static/images/{location}/pm10_grouped.png', width=Inches(6))

    document.add_heading('O3 Forecast', level=2)
    document.add_picture(f'static/images/{location}/o3_forecast.png', width=Inches(6))

    document.add_heading('O3 Forecast (Grouped Bar Chart)', level=2)
    document.add_picture(f'static/images/{location}/o3_grouped.png', width=Inches(6))

    word_file = f'report_{location}.docx'
    document.save(word_file)

    return send_file(word_file, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
